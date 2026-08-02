from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
OUTPUT_PATH = OUTPUT_DIR / "学生事务管家-运行与数据存储详细说明.docx"
DIAGRAM_PATH = OUTPUT_DIR / ".system-architecture.png"

BLUE = "2E74B5"
DARK_BLUE = "0B2545"
DEEP_GREEN = "1F6B5B"
MUTED = "5E6B73"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E8F2EF"
LIGHT_GRAY = "F4F6F9"
BORDER = "CDD6DF"
CAUTION = "B66A32"
WHITE = "FFFFFF"
BLACK = "20252A"

FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft YaHei"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must total {CONTENT_DXA}, got {sum(widths)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def mark_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 name: str = FONT_LATIN) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_CJK)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def apply_style_font(style, size: float, color: str = BLACK, bold: bool = False) -> None:
    style.font.name = FONT_LATIN
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_CJK)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    apply_style_font(normal, 11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        apply_style_font(style, size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        apply_style_font(style, 11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "学生事务管家  |  运行与数据存储说明"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    for run in hp.runs:
        set_run_font(run, size=9, color=MUTED)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(fp, "PAGE")
    run = fp.add_run(" 页  /  共 ")
    set_run_font(run, size=9, color=MUTED)
    add_page_field(fp, "NUMPAGES")
    run = fp.add_run(" 页")
    set_run_font(run, size=9, color=MUTED)


def add_paragraph(doc, text: str, *, bold_lead: str | None = None,
                  color: str = BLACK, after: float = 6,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=11, color=color, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=11, color=color)
    else:
        run = p.add_run(text)
        set_run_font(run, size=11, color=color)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run, size=11)


def add_numbered_list(doc, items: tuple[str, ...]) -> None:
    """Add a real numbered list that always restarts at 1."""
    style = doc.styles["List Number"]
    style_num_id = style.element.pPr.numPr.numId.val
    numbering = doc.part.numbering_part.element
    abstract_num_id = numbering.num_having_numId(style_num_id).abstractNumId.val
    num = numbering.add_num(abstract_num_id)
    num.add_lvlOverride(ilvl=0).add_startOverride(1)
    num_id = num.numId

    for text in items:
        p = doc.add_paragraph(style="List Number")
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        num_pr.get_or_add_numId().val = num_id
        run = p.add_run(text)
        set_run_font(run, size=11)


def add_heading(doc, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level],
                 color=BLUE if level < 3 else DARK_BLUE, bold=True)


def add_callout(doc, label: str, text: str, *, fill: str = LIGHT_GRAY,
                accent: str = DEEP_GREEN) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    lead = p.add_run(f"{label}  ")
    set_run_font(lead, size=10.5, color=accent, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.5, color=BLACK)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    mark_repeat_header(table.rows[0])

    for row_data in rows:
        row = table.add_row()
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        cells = row.cells
        for index, value in enumerate(row_data):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if index >= 2 and len(headers) >= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, size=9.2, color=BLACK, bold=(index == 0))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def make_architecture_diagram(path: Path) -> None:
    width, height = 1600, 880
    image = Image.new("RGB", (width, height), "#F7F9FA")
    draw = ImageDraw.Draw(image)
    font_path = Path("C:/Windows/Fonts/msyh.ttc")
    fallback = Path("C:/Windows/Fonts/arial.ttf")
    chosen = font_path if font_path.exists() else fallback
    title_font = ImageFont.truetype(str(chosen), 42)
    box_title = ImageFont.truetype(str(chosen), 31)
    body_font = ImageFont.truetype(str(chosen), 24)
    small_font = ImageFont.truetype(str(chosen), 21)

    draw.text((70, 45), "学生事务管家的实际运行与数据流", font=title_font, fill="#0B2545")

    def rounded_box(xy, fill, outline, title, lines):
        draw.rounded_rectangle(xy, radius=28, fill=fill, outline=outline, width=4)
        x1, y1, x2, _ = xy
        draw.text((x1 + 34, y1 + 30), title, font=box_title, fill="#0B2545")
        y = y1 + 92
        for line in lines:
            draw.text((x1 + 34, y), line, font=body_font, fill="#37474F")
            y += 42

    rounded_box((65, 175, 485, 485), "#FFFFFF", "#2E74B5", "浏览器中的 React 网页", [
        "页面、任务编辑、日历", "本机文件文字读取", "用户确认与修改",
    ])
    rounded_box((590, 175, 1010, 485), "#FFFFFF", "#1F6B5B", "Cloudflare Worker", [
        "托管网页静态资源", "处理 /api/deepseek", "服务端读取 Secret",
    ])
    rounded_box((1115, 175, 1535, 485), "#FFFFFF", "#B66A32", "DeepSeek V4 Flash", [
        "结构化任务建议", "基于引用的问答", "不直接读取浏览器数据库",
    ])
    rounded_box((65, 610, 485, 805), "#E8EEF5", "#2E74B5", "IndexedDB（本机）", [
        "保存任务、来源、草稿", "项目、课程、历史、提醒",
    ])
    rounded_box((590, 610, 1010, 805), "#E8F2EF", "#1F6B5B", "Cloudflare Secret", [
        "只保存 DEEPSEEK_API_KEY", "网页 JavaScript 无法读取",
    ])

    def arrow(start, end, color):
        draw.line([start, end], fill=color, width=7)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 18, ey - 12), (ex - 18, ey + 12)], fill=color)

    arrow((485, 330), (590, 330), "#2E74B5")
    arrow((1010, 330), (1115, 330), "#B66A32")
    arrow((275, 485), (275, 610), "#2E74B5")
    arrow((800, 610), (800, 485), "#1F6B5B")
    draw.text((495, 285), "网页与同源 API", font=small_font, fill="#5E6B73")
    draw.text((1015, 285), "受控文本请求", font=small_font, fill="#5E6B73")
    draw.text((295, 525), "自动持久化", font=small_font, fill="#5E6B73")
    draw.text((820, 525), "密钥绑定", font=small_font, fill="#5E6B73")
    image.save(path, "PNG")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("产品技术说明书")
    set_run_font(run, size=11, color=CAUTION, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("学生事务管家")
    set_run_font(run, size=30, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run("网页运行机制、数据存储、隐私边界与备份恢复")
    set_run_font(run, size=15, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(120)
    run = p.add_run("适用于当前 Cloudflare 生产版本")
    set_run_font(run, size=10.5, color=MUTED, italic=True)

    for label, value in (
        ("生产网址", "https://student-affairs.site"),
        ("部署方式", "Cloudflare Worker + Static Assets"),
        ("智能模型", "DeepSeek V4 Flash（服务端代理）"),
        ("文档日期", "2026 年 8 月 3 日"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        lead = p.add_run(f"{label}：")
        set_run_font(lead, size=10.5, color=DARK_BLUE, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(48)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("核心结论：网页运行在云端，业务数据默认保存在当前浏览器本机。")
    set_run_font(run, size=11, color=DEEP_GREEN, bold=True)
    p.add_run().add_break(WD_BREAK.PAGE)


def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    make_architecture_diagram(DIAGRAM_PATH)

    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "学生事务管家：运行与数据存储详细说明"
    doc.core_properties.subject = "Cloudflare 部署、IndexedDB 数据存储与 DeepSeek 服务端代理说明"
    doc.core_properties.author = "学生事务管家项目"
    doc.core_properties.keywords = "学生事务管家, Cloudflare, IndexedDB, DeepSeek, 数据存储, 隐私"
    add_cover(doc)

    add_heading(doc, "阅读导览", 1)
    add_paragraph(doc, "这份说明面向日常使用者、维护者和未来开发者。前半部分回答“网页怎么运行、数据在哪里”；后半部分说明 DeepSeek、文件上传、提醒、备份和可选同步服务的真实边界。")
    add_callout(doc, "最重要的结论", "Cloudflare 托管网页和 DeepSeek 代理，但任务、来源、待确认草稿、项目、材料勾选、提醒设置和修改历史默认都保存在当前浏览器的 IndexedDB 中。目前没有账号系统，也没有生产级跨设备云同步。")
    for item in (
        "第 1—3 章：系统如何加载、请求如何流转、为什么 Cloudflare 控制台显示 Worker。",
        "第 4—7 章：本机数据库的精确位置、保存内容、文件与 DeepSeek 的数据边界。",
        "第 8—11 章：提醒、备份恢复、数据丢失风险与可选本机同步服务。",
        "第 12—14 章：如何自行检查数据、隐私与安全建议、常见问题。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "1. 一句话理解整个系统", 1)
    add_paragraph(doc, "学生事务管家是一套运行在浏览器中的 React 单页应用。Cloudflare 负责把网页文件送到浏览器，并在你主动使用智能能力时代理 DeepSeek；浏览器负责界面交互、文件本机读取和业务数据持久化。")
    add_table(doc,
              ["组成部分", "主要职责", "是否保存业务数据", "当前状态"],
              [
                  ["浏览器 React 网页", "展示页面、编辑任务、读取本机文件文字、管理确认流程", "是，保存到 IndexedDB", "真实可用"],
                  ["Cloudflare Worker", "托管网页、SPA 回退、同源 DeepSeek API 代理", "否，没有 D1/KV/R2 业务库", "真实可用"],
                  ["Cloudflare Secret", "保存 DEEPSEEK_API_KEY", "仅保存服务密钥", "真实可用"],
                  ["DeepSeek V4 Flash", "生成结构化任务建议与受引用约束的回答", "由服务商按其政策处理请求", "真实可用"],
                  ["可选本机 Node 服务", "手动同步、邮件队列和受控网页读取基础", "启用时写入本机 .data", "默认未启动"],
              ],
              [2100, 3460, 2000, 1800])

    add_heading(doc, "2. 系统架构与数据流", 1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(DIAGRAM_PATH), width=Inches(6.35))
    doc_pr = run._r.xpath(".//wp:docPr")[0]
    doc_pr.set("title", "学生事务管家系统架构")
    doc_pr.set("descr", "浏览器 React 网页、IndexedDB、Cloudflare Worker、Cloudflare Secret 与 DeepSeek V4 Flash 之间的数据流和隔离关系。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("图 1  当前生产环境的真实架构")
    set_run_font(run, size=9, color=MUTED, italic=True)
    add_paragraph(doc, "图中的关键隔离是：IndexedDB 位于用户浏览器中；Cloudflare Worker 不会直接读取整个 IndexedDB。只有用户主动触发智能整理或云端问答时，浏览器才会把当次允许的文本发送到同源 API。")

    add_heading(doc, "3. 打开网页后发生什么", 1)
    add_numbered_list(doc, (
        "浏览器访问 https://student-affairs.site，DNS 将请求交给 Cloudflare。",
        "Cloudflare Worker 返回 Vite 构建的 HTML、JavaScript、CSS 和 PDF 解析组件。",
        "React 在浏览器中启动，先显示应用外壳和“正在恢复本机工作区”状态。",
        "应用打开 IndexedDB 数据库 student-affairs-steward，读取 workspace 存储区中键为 current 的记录。",
        "如果找到有效工作区，就恢复任务、来源、草稿、项目、课程、接入状态和知识设置。",
        "如果没有 IndexedDB 记录，则尝试读取旧 localStorage 数据进行迁移；仍没有数据时才使用演示初始数据。",
        "完成水合后，任何工作区变化都会触发自动保存。刷新和关闭浏览器后再次打开，数据由 IndexedDB 恢复。",
    ))
    add_callout(doc, "为什么控制台显示 Worker", "Cloudflare Worker 同时承担静态网站入口和服务端 API，因此控制台把项目归类为 Worker。对用户来说，它依然是完整网页；“Worker”描述的是部署运行方式，不代表页面不是网站。", fill=LIGHT_BLUE, accent=BLUE)

    add_heading(doc, "4. 数据究竟存在哪里", 1)
    add_heading(doc, "4.1 IndexedDB 主工作区", 2)
    add_table(doc,
              ["项目", "实际值"],
              [
                  ["数据库名称", "student-affairs-steward"],
                  ["数据库版本", "1（内部工作区 schema 为 v5）"],
                  ["Object Store", "workspace"],
                  ["记录键", "current"],
                  ["所属 Origin", "https://student-affairs.site"],
                  ["保存方式", "React 状态变化后自动覆盖同一条版本化工作区记录"],
              ],
              [2600, 6760])
    add_paragraph(doc, "IndexedDB 是浏览器提供的结构化本地数据库。它比 localStorage 更适合保存任务、数组、历史记录和较大的来源文字。数据位于浏览器用户配置目录中，不是项目 Git 文件，也不是 Cloudflare 数据库文件。")

    add_heading(doc, "4.2 Origin 隔离", 2)
    add_paragraph(doc, "浏览器以“协议 + 域名 + 端口”划分站点存储。因此下列地址拥有彼此独立的数据空间：")
    for item in (
        "https://student-affairs.site（当前生产主站）",
        "https://student-affairs-manager.nightsdell.workers.dev（Cloudflare 备用地址）",
        "Firebase 旧地址",
        "http://localhost:5173 或其他本地开发地址",
    ):
        add_bullet(doc, item)
    add_callout(doc, "容易误解的地方", "即使四个地址展示同一套页面代码，它们的 IndexedDB 也互不相通。换域名访问时看不到原来的任务，不等于数据被服务器删除，而是进入了另一个浏览器存储空间。", fill="FFF7EE", accent=CAUTION)

    add_heading(doc, "4.3 localStorage 的剩余用途", 2)
    add_paragraph(doc, "核心业务数据已经迁移到 IndexedDB。localStorage 主要保存新手教程是否完成，并兼容读取旧版工作区、任务和来源。")
    add_table(doc,
              ["键名", "用途", "是否核心数据"],
              [
                  ["student-affairs-steward:onboarding:v1", "记录教程已完成", "否，界面偏好"],
                  ["student-affairs-steward:workspace:v2", "旧版工作区迁移来源", "旧数据兼容"],
                  ["student-affairs-steward:tasks:v1", "旧版任务迁移来源", "旧数据兼容"],
                  ["student-affairs-steward:sources:v1", "旧版来源迁移来源", "旧数据兼容"],
              ],
              [3600, 3900, 1860])

    add_heading(doc, "5. IndexedDB 保存哪些内容", 1)
    add_table(doc,
              ["实体", "保存的典型内容"],
              [
                  ["任务 Tasks", "标题、分类、状态、截止时间、预计耗时、下一步、描述、优先级与理由"],
                  ["材料 Materials", "材料名称、是否完成、所属任务/项目/来源关联"],
                  ["提醒 Reminders", "渠道、计划时间、是否启用"],
                  ["修改历史 History", "修改字段、修改前后值、时间、用户或系统操作来源"],
                  ["来源 Sources", "来源类型、标题、摘要、非链接来源的提取文字、网址、提取方式和重复提示"],
                  ["待确认草稿 Drafts", "一份来源对应的多条任务建议、逐条确认/拒绝状态与更新时间"],
                  ["项目 Projects", "项目标题、任务关联、来源关联和里程碑"],
                  ["课程 Course blocks", "星期、开始时间、结束时间，用于给出避让建议"],
                  ["网页监测记录", "用户授权的网址、手工粘贴正文基线、哈希、差异摘要与检查方法"],
                  ["接入状态", "本机服务地址、同步修订、网页监测、微信/跨设备接入意向"],
                  ["知识设置", "本地知识检索授权时间"],
              ],
              [2600, 6760])
    add_callout(doc, "隐私提示", "来源正文可能包含老师通知、课程安排或个人材料中的文字。虽然文件本体不会保存，但提取出的文字会随来源进入本机工作区，并包含在完整 JSON 备份中。")

    add_heading(doc, "6. 文件上传时发生什么", 1)
    add_heading(doc, "6.1 TXT 与 Markdown", 2)
    add_paragraph(doc, "文件通过浏览器 File API 在本机读取。系统保存清洗后的文字，不保存原始文件二进制，也不会仅因为选择文件就自动上传到服务器。")
    add_heading(doc, "6.2 带文本层的 PDF", 2)
    add_paragraph(doc, "PDF.js 在浏览器中读取 PDF 文本层，最多处理前 80 页；提取文字规范化后最多保留约 50,000 个字符。加密、损坏或无文本层 PDF 会显示真实失败状态。")
    add_heading(doc, "6.3 图片和扫描 PDF", 2)
    add_paragraph(doc, "当前没有生产 OCR。图片文件本体不会写入 IndexedDB，也不会因为上传动作自动发送给 DeepSeek。用户必须人工补充截图或扫描件中的原文。")
    add_table(doc,
              ["输入类型", "本机处理", "保存内容", "自动发送云端"],
              [
                  ["TXT / Markdown", "浏览器读取文字", "提取文字与摘要", "否；需主动智能整理"],
                  ["文本层 PDF", "PDF.js 解析前 80 页", "提取文字与摘要", "否；需主动智能整理"],
                  ["图片", "仅选择文件", "人工补充的文字", "否"],
                  ["扫描 PDF", "检测不到文本层", "人工补充的文字", "否"],
                  ["网页链接", "默认仅保存链接", "网址和摘要", "不自动抓取或发送"],
              ],
              [2000, 2600, 2500, 2260])

    add_heading(doc, "7. DeepSeek 如何工作", 1)
    add_heading(doc, "7.1 服务端代理", 2)
    add_paragraph(doc, "网页不会持有 DeepSeek API Key。浏览器把允许的内容发送到同源 /api/deepseek 或 /api/deepseek/extract；Cloudflare Worker 从 DEEPSEEK_API_KEY Secret 中读取密钥，再访问 DeepSeek API。")
    add_paragraph(doc, "Worker 当前固定使用 deepseek-v4-flash，并关闭思考模式。服务端限制请求来源、正文大小、任务数量和基础请求频率；密钥不写入浏览器、IndexedDB、导出文件、URL、日志或 Git。")

    add_heading(doc, "7.2 智能整理发送范围", 2)
    add_table(doc,
              ["会发送", "不会发送"],
              [
                  ["当前来源类型和标题", "DeepSeek API Key"],
                  ["当前输入或浏览器提取的文字（最多 24,000 字符）", "文件二进制、图片本体"],
                  ["参考时间和浏览器时区", "整个 IndexedDB 或完整工作区"],
                  ["用户主动触发的当次请求", "未主动读取的网页正文"],
              ],
              [4680, 4680])
    add_paragraph(doc, "DeepSeek 返回的是“建议草稿”。系统要求标题、截止时间、预计耗时、下一步、材料、证据等结构满足校验，并最多返回 20 条。建议先进入待确认队列；只有用户逐条确认、部分确认或全部确认后，才创建正式任务。")

    add_heading(doc, "7.3 知识问答发送范围", 2)
    add_paragraph(doc, "本地知识检索必须先授权。每次点击“使用 DeepSeek 回答”都会再次显示发送确认。本次只发送问题和本地命中的最多 4 条引用摘要；不会把整个工作区发送给模型。无本地命中时，系统不应凭模型常识补写。")
    add_callout(doc, "服务商边界", "本项目代码不会把请求正文写入 Cloudflare 数据库，也没有主动记录正文日志；但文本会经过 Cloudflare 网络并由 DeepSeek 处理，因此仍应遵守 DeepSeek 的服务条款和隐私政策，不建议发送不必要的身份证号、账号密码、医疗信息或完整敏感材料。", fill="FFF7EE", accent=CAUTION)

    add_heading(doc, "8. 提醒是怎样保存和触发的", 1)
    add_paragraph(doc, "提醒规则本身保存在任务数据中，因此刷新后仍可恢复。浏览器通知权限由浏览器管理，不属于 IndexedDB 工作区。")
    add_bullet(doc, "只有用户点击并允许通知后，网页才能创建浏览器通知。")
    add_bullet(doc, "当前使用页面内计时器；页面关闭后没有 Service Worker 在后台继续调度。")
    add_bullet(doc, "已触发通知的内存去重集合会在刷新后重置。")
    add_bullet(doc, "邮件与微信目前没有生产发送通道；保存计划不等于已经发送。")

    add_heading(doc, "9. 备份、恢复与清空", 1)
    add_heading(doc, "9.1 JSON 完整备份", 2)
    add_paragraph(doc, "“项目档案”页面可导出 student-affairs-backup-YYYY-MM-DD.json。该文件是完整结构化快照，可用于恢复任务、来源、草稿、项目、课程、接入记录和知识设置。")
    add_callout(doc, "备份同样敏感", "JSON 可能包含来源全文、任务历史和材料信息，应像私人文档一样保存。不要公开上传到 Git、论坛或无访问控制的网盘。")
    add_heading(doc, "9.2 导入与替换", 2)
    add_paragraph(doc, "导入时会先解析并迁移 schema v3/v4/v5 数据。有效备份会替换当前浏览器内的工作区；无法识别的 JSON 会被拒绝。导入前建议先导出当前工作区，防止误覆盖。")
    add_heading(doc, "9.3 清空", 2)
    add_paragraph(doc, "清空操作需要在应用内连续两次确认。确认后会删除 IndexedDB 中键为 current 的工作区记录，并清空页面状态；该操作无法撤销，只能通过之前导出的 JSON 恢复。")

    add_heading(doc, "10. 什么情况下数据会消失", 1)
    add_table(doc,
              ["情况", "结果", "建议"],
              [
                  ["刷新或正常关闭浏览器", "通常可恢复", "无需处理"],
                  ["清除 student-affairs.site 站点数据", "本机工作区被删除", "清除前导出 JSON"],
                  ["清除全部浏览器数据或重装并删除配置", "可能全部丢失", "定期离线备份"],
                  ["使用无痕/隐私窗口", "关闭窗口后可能删除", "不要长期在无痕模式使用"],
                  ["换电脑、换浏览器、换浏览器用户", "不会自动出现原数据", "导出并导入 JSON"],
                  ["改用 workers.dev、Firebase 或 localhost", "进入独立存储空间", "坚持使用生产主域名"],
                  ["浏览器存储损坏或配额被回收", "存在小概率丢失风险", "保留周期性备份"],
              ],
              [2800, 2500, 4060])

    add_heading(doc, "11. 可选本机同步服务", 1)
    add_paragraph(doc, "仓库包含一个可选 Node.js 服务，默认只监听 127.0.0.1:8787。它目前没有部署到 Cloudflare，也不等于云账号同步。")
    add_bullet(doc, "启用时需要在 .env 配置至少 20 位 SAM_SYNC_TOKEN。")
    add_bullet(doc, "工作区保存到项目目录 .data/workspace.json。")
    add_bullet(doc, "文件以临时文件写入后重命名，降低写入中断造成的损坏。")
    add_bullet(doc, "工作区使用内容修订号检测冲突；发生冲突时由用户选择拉取或确认覆盖。")
    add_bullet(doc, "前端令牌只保存在当前 React 页面内存中，刷新后清空。")
    add_bullet(doc, "当前没有账号登录、自动后台同步、设备管理、端到端加密或互联网托管。")

    add_heading(doc, "12. 如何自己查看本机数据", 1)
    add_numbered_list(doc, (
        "在 Chrome 或 Edge 中打开 https://student-affairs.site。",
        "按 F12 打开开发者工具。",
        "选择“应用”或 Application 面板。",
        "在左侧展开 IndexedDB。",
        "展开 student-affairs-steward，再选择 workspace。",
        "查看键为 current 的记录；其中就是完整工作区对象。",
        "如需查看教程偏好，展开 Local Storage 并选择 student-affairs.site。",
    ))
    add_callout(doc, "不要直接编辑", "开发者工具适合检查，不适合作为日常编辑入口。手工修改错误字段可能使工作区无法通过 schema 校验。需要调整数据请使用网页界面；需要迁移请使用 JSON 导出/导入。", fill="FFF7EE", accent=CAUTION)

    add_heading(doc, "13. 隐私与安全建议", 1)
    for item in (
        "始终从 https://student-affairs.site 使用正式版本，避免数据散落到多个域名。",
        "每周或每次重要项目更新后导出一次 JSON 备份。",
        "发送 DeepSeek 前删去与任务识别无关的身份证号、密码、手机号和完整个人材料。",
        "不要把 JSON 备份、.env、同步令牌或 API Key 提交到 Git。",
        "多人共用电脑时使用独立浏览器用户，并启用系统账户密码。",
        "不再使用的 DeepSeek Key 应在 DeepSeek 控制台撤销；Cloudflare Secret 中只保留当前有效 Key。",
        "清除浏览器数据、重装浏览器或更换设备前先导出备份。",
    ):
        add_bullet(doc, item)

    add_heading(doc, "14. 常见问题", 1)
    faq = [
        ("为什么刷新后任务还在？", "因为 React 会在工作区变化后写入 IndexedDB，重新打开时再从同一 Origin 的 IndexedDB 恢复。"),
        ("为什么换浏览器后任务没了？", "不是任务被服务器删除，而是新浏览器拥有独立的站点存储。使用 JSON 导出和导入完成迁移。"),
        ("Cloudflare 会保存我的全部任务吗？", "当前不会。Worker 没有 D1、KV 或 R2 业务数据库绑定，只托管静态网页并代理 DeepSeek。"),
        ("DeepSeek 能看到我所有历史任务吗？", "智能整理只接收当次输入文字；知识问答只接收问题和最多 4 条命中摘要，并要求逐次确认。"),
        ("上传 PDF 是否等于上传到云端？", "不是。PDF 文本层首先在浏览器本机解析；只有主动点击智能整理，提取出的当前文字才会发送给服务端代理。"),
        ("图片为什么无法自动识别？", "生产 OCR 尚未接通。系统诚实要求人工补充文字，不会把图片假装成已识别。"),
        ("关掉网页后还会提醒吗？", "当前不会可靠后台提醒。提醒规则仍在，但页面内计时器停止；需要后续 Service Worker 或服务端推送能力。"),
        ("怎样获得真正跨设备同步？", "需要账号认证、受信任后端、设备管理、传输加密、冲突策略和备份恢复。目前这些能力尚未上线。"),
    ]
    for question, answer in faq:
        add_heading(doc, question, 2)
        add_paragraph(doc, answer)

    add_heading(doc, "附录 A：当前代码位置速查", 1)
    add_table(doc,
              ["功能", "文件"],
              [
                  ["Cloudflare 路由与静态资源", "wrangler.jsonc"],
                  ["Worker 与 DeepSeek 代理", "cloudflare/worker.mjs"],
                  ["IndexedDB Repository", "src/lib/repository.ts"],
                  ["旧 localStorage 迁移", "src/lib/storage.ts"],
                  ["应用加载与自动保存", "src/App.tsx"],
                  ["工作区数据类型", "src/types.ts"],
                  ["文件本机解析", "src/lib/fileExtraction.ts"],
                  ["智能整理客户端", "src/lib/deepseekExtraction.ts"],
                  ["知识问答客户端", "src/lib/deepseek.ts"],
                  ["JSON 备份与清空", "src/components/WorkspaceControls.tsx"],
                  ["可选本机文件存储", "server/workspace-store.mjs"],
              ],
              [3100, 6260])

    add_heading(doc, "附录 B：当前能力状态", 1)
    add_table(doc,
              ["能力", "状态", "说明"],
              [
                  ["Cloudflare 网页部署", "已接通", "生产主域名 student-affairs.site"],
                  ["IndexedDB 本机持久化", "已接通", "刷新和关闭后可恢复"],
                  ["DeepSeek V4 Flash", "已接通", "Cloudflare Secret + 同源代理"],
                  ["TXT/Markdown/PDF 文本层", "已接通", "浏览器本机解析"],
                  ["图片/扫描件 OCR", "未接通", "需要人工补充文字"],
                  ["浏览器前台提醒", "已接通", "仅页面存活期间调度"],
                  ["真实邮件发送", "未接通", "仅有服务端接口与队列基础"],
                  ["微信授权与消息读取", "未接通", "等待官方平台审批与合规后端"],
                  ["跨设备云同步", "未接通", "没有生产账号和托管后端"],
                  ["Obsidian 导出", "已接通", "静态 Markdown/ZIP 或授权文件夹写入"],
              ],
              [2700, 1700, 4960])

    add_callout(doc, "维护原则", "界面中显示“已连接”“已发送”“已同步”等状态时，必须有真实服务端证据。任何自动分类、日期、耗时、优先级和任务拆分都只是建议，最终由用户确认，确认后的内容不得被后续解析静默覆盖。", fill=LIGHT_GREEN, accent=DEEP_GREEN)

    doc.save(OUTPUT_PATH)
    DIAGRAM_PATH.unlink(missing_ok=True)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_document()
    document = Document(output)
    print(f"created={output}")
    print(f"paragraphs={len(document.paragraphs)} tables={len(document.tables)} sections={len(document.sections)}")
